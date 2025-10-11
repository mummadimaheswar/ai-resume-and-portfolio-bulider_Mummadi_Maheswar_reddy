"""
File processing utilities for various formats
"""

import os
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class FileProcessor:
    
    def read_file(self, file_path):
        """Read content from various file formats"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.txt':
            return self._read_txt(file_path)
        elif file_ext == '.pdf':
            return self._read_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return self._read_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def write_file(self, filename, content, format_type):
        """Write content to specified format"""
        if format_type == 'txt':
            self._write_txt(filename, content)
        elif format_type == 'docx':
            self._write_docx(filename, content)
        else:
            # Default to txt for other formats
            self._write_txt(filename.replace(f'.{format_type}', '.txt'), content)
    
    def _read_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _read_pdf(self, file_path):
        content = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        return content
    
    def _read_docx(self, file_path):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content
    
    def _write_txt(self, filename, content):
        with open(filename, 'w', encoding='utf-8') as file:
            file.write(content)
    
    def _write_docx(self, filename, content):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        doc = Document()
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        doc.save(filename)